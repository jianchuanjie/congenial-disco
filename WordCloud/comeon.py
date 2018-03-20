#!python2
#coding=utf-8
import matplotlib.pyplot as plt
from wordcloud import WordCloud,ImageColorGenerator
import jieba
from scipy.misc import imread
from PIL import Image
from PIL import ImageEnhance,ImageOps
import docx
import sys
import io
import codecs
import requests
from bs4 import BeautifulSoup
import  regex as re
import chardet
import os
reload(sys)
sys.setdefaultencoding('utf-8')
f=open("zimu.txt",'r')
img_mask=imread("tupian2.jpg")


def increaseContrast(tupian="t1.jpg"):
	image = Image.open(tupian)
	enh_con = ImageEnhance.Contrast(image)  
	contrast = 1.5  
	image_contrasted = enh_con.enhance(contrast)
	image_contrasted.save("cons.jpg")
	return(imread("cons.jpg"))  
	#image_contrasted.show()  
#tupian="tupian2.jpg"
#changeToGray(tupian)
def setBackgroundColor(color="white"):
	return color

def setFontPath(path='16.ttf'):
	return path

def setColorMap(mp="spring"):
	return mp

def readDocument():
    '''
    获取文档对象，将文档内容按段落读入，并存入doc中
    '''
    file = docx.Document("dang.docx")
    doc = ""
    for para in file.paragraphs:
        doc = doc + para.text
    return doc

def segment(fil=f):
    '''
    用jieba分词对输入文档进行分词，并保存至本地（根据情况可跳过）
    '''
    doc=""
    for line in fil:
    	doc=doc+line
    seg_list = " ".join(jieba.cut(doc, cut_all=False)) #seg_list为str类型
    document_after_segment = open('分词结果1.txt'.decode('utf-8'),  'w+')
    document_after_segment.write(seg_list)
    document_after_segment.close()
    return seg_list
    
def removeStopWords(seg_list):
    '''
    自行下载stopwords1893.txt停用词表，该函数实现去停用词,英文本身有去停用词的功能，所以不必调用
    此处编码耗时太久，终于解决！各种编码类型转换为Unicode，然后关键是要去掉多余的空格和\r，输出中文文件名需要从python默认的utf-8编码解码到Unicode，然后系统自行编码即可
    '''
    wordlist_stopwords_removed = []
    stop_words = codecs.open('stops1.txt')
    stop_words_text = stop_words.read().decode('utf-8')
    p=[]
    stop_words.close()
    stop_words_text_list = stop_words_text .split('\n')
    after_seg_text_list = seg_list.split(' ')
    for l in stop_words_text_list:
    	p.append(l.replace('\r','').replace(' ',''))
    wordlist_stopwords_removed=[val for val in after_seg_text_list if val not in p]
    without_stopwords = open('分词结果(去停用词)1.txt'.decode('utf-8'), 'w')
    without_stopwords.write(' '.join(wordlist_stopwords_removed))
    return ' '.join(wordlist_stopwords_removed)

def wordCount(segment_list):
    '''
        该函数实现词频的统计，并将统计结果存储至本地。
        在制作词云的过程中用不到，主要是在画词频统计图时用到。
    '''
    word_lst = []
    word_dict = {}
    with open('词频统计(去停用词).txt'.decode('utf-8'),'w') as wf2: 
        word_lst.append(segment_list.split(' ')) 
        for item in word_lst:
            for item2 in item:
                if item2 not in word_dict: 
                    word_dict[item2] = 1
                else:
                    word_dict[item2] += 1
        lis=[]
        lis=sorted(word_dict.items(), \
        key = lambda item:item[1], reverse=True)
        for key in lis:#如果更改词频，考虑从此处入手
            wf2.write(key[0]+' '+str(key[1])+'\n') 
    wf2.close()
    return word_dict
def changeFre(word_dict,changelist):
    with open('词频统计(去停用词).txt'.decode('utf-8'),'w') as wf2:
        for item in changelist:
            word_dict[item[0].decode('utf-8')] = item[1]
        lis=[]
        lis=sorted(word_dict.items(), \
        key = lambda item:item[1], reverse=True)
        for key in lis:#如果更改词频，考虑从此处入手
            wf2.write(key[0]+' '+str(key[1])+'\n') 
    wf2.close()
    return word_dict
        

#segment(readDocument())
def WCcreate(seg_list,background_color=setBackgroundColor(), #背景颜色
    max_words=100,# 词云显示的最大词数
    font_path=setFontPath(),
    mask=img_mask,#设置背景图片
    max_font_size=100, #字体最大值
    random_state=42,
    relative_scaling=0.4,
    colormap=setColorMap(),
    ):
    
    '''
        制作词云
        设置词云参数
    '''
    color_mask = mask # 读取背景图片,注意路径
    wc = WordCloud(
        #设置字体，不指定就会出现乱码，注意字体路径
        font_path=font_path,
        #font_path=path.join(d,'simsun.ttc'),
        #设置背景色
        background_color=background_color,
        #词云形状
        mask=color_mask,
        #允许最大词汇
        max_words=200,
        #最大号字体
        max_font_size=60,
        colormap=colormap
    )
    wc.fit_words(seg_list) # 产生词云
    image_colors = ImageColorGenerator(color_mask)
    wc.to_file("E:/ciyun.jpg") #保存图片
    #  显示词云图片
    '''plt.imshow(wc, interpolation="bilinear")
    plt.axis('off')

    #这里主要为了实现词云图片按照图片颜色取色
    plt.figure()
    plt.imshow(wc.recolor(color_func=image_colors), interpolation="bilinear")
    plt.axis("off")

    plt.show()'''

def readURL(u='http://www.baidu.com'):
    try:
        r=requests.get(u,timeout=300)
        print(r.status_code)
        r.encoding='r.encoding'
		
    except:
		return "error"
    #f=re.compile('''\\p{han}+''')
    soup=BeautifulSoup(r.text ,"lxml")
    s=''
    for child in soup.body.children:
    	s=s+str(child).decode('utf-8')#.encode('gb18030')
    fe=re.findall(r"([\u4e00-\u9fa5]+)",s)
    return ' '.join(fe)


if __name__ == "__main__":
    #doc = readDocument()
    #segment_list = segment(doc)
    segment_list = readURL('https://www.jianshu.com/p/d72a7b1fe807')
    segment_list_remove_stopwords = removeStopWords(segment_list)
    WCcreate(wordCount(segment_list_remove_stopwords))
    WCcreate(changeFre(wordCount(segment_list_remove_stopwords),[['我来啦',10]]))
    #readURL()
    #changeToGray()