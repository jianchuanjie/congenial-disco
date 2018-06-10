#!/usr/bin/env python3

import matplotlib.pyplot as plt
from wordcloud import WordCloud, ImageColorGenerator
import jieba
from scipy.misc import imread
from PIL import Image
import docx
import regex as re
import os, collections


# f = open('zimu.txt')
wd = os.path.join(os.getcwd(), 'WC')
FONTS_PATH = os.path.join(os.getcwd(), 'fonts')

img_mask_path = os.path.join(wd, "tupian2.jpg")
doc_path_default = os.path.join(wd, "dang.docx")


def changeToGray(tupian="tupian2.jpg"):
    # 将彩色图片转换成黑白图片
    im=Image.open(tupian).convert('L')
    # 保存图片
    return im


#tupian="tupian2.jpg"
#changeToGray(tupian)
def setBackgroundColor(color="white"):
    return color


def setFontPath(path=os.path.join(FONTS_PATH, 'fan' + '.ttf')):
    return path


def setColorMap(mp="spring"):
    return mp


def setSegList(doc_path=doc_path_default):
    doc = readDocument(doc_path)
    segment_list = segment(doc)
    seg_list = removeStopWords(segment_list)
    return seg_list

def readDocument(doc_path=doc_path_default):
    '''
    获取文档对象，将文档内容按段落读入，并存入doc中
    '''
    file = docx.Document(doc_path)
    doc = ""
    for para in file.paragraphs:
        doc = doc + para.text
    # print(type(doc))
    return doc


def segment(fil):
    '''
    用jieba分词对输入文档进行分词，并保存至本地（根据情况可跳过）
    '''
    doc=""
    for line in fil:
        doc=doc+line
    seg_list = " ".join(jieba.cut(doc, cut_all=False)) #seg_list为str类型

    # document_after_segment = open('分词结果.txt', 'w+')
    # document_after_segment.write(seg_list)
    # document_after_segment.close()
    return removeStopWords(seg_list)


def removeStopWords(seg_list):
    '''
    自行下载stopwords1893.txt停用词表，该函数实现去停用词,英文本身有去停用词的功能，所以不必调用
    '''
    wordlist_stopwords_removed = []

    stop_words = open(os.path.join(wd, 'stops1.txt'),encoding = 'utf-8')
    stop_words_text = stop_words.read()

    stop_words.close()

    stop_words_text_list = stop_words_text.split('\n')
    after_seg_text_list = seg_list.split(' ')
    wordlist_stopwords_removed = [val for val in after_seg_text_list \
        if val not in str(stop_words_text_list)]
    print(type(stop_words_text_list[0]))
    # without_stopwords = open('分词结果(去停用词).txt', 'w')
    # without_stopwords.write(' '.join(wordlist_stopwords_removed))
    return ' '.join(wordlist_stopwords_removed)


def wordCount(segment_list):  #函数更新
    '''
        该函数实现词频的统计，并将统计结果存储至本地。
        在制作词云的过程中用不到，主要是在画词频统计图时用到。
    '''
    word_lst = []
    word_dict = {}
#    with open('词频统计(去停用词).txt','w') as wf2: 
    word_lst.append(segment_list.split(' ')) 

    for item in word_lst:
            for item2 in item:
                if re.match(r"\S.*",str(item2)):

                  if item2 not in word_dict: 

                      word_dict[item2] = 1
                  else:
                      word_dict[item2] += 1

    word_dict_sorted = collections.OrderedDict(sorted(word_dict.items(), \
    key = lambda item:item[1], reverse=True))#按照词频从大到小排序
#        for key in word_dict_sorted:#如果更改词频，考虑从此处入手
#            wf2.write(key+' '+str(word_dict_sorted[key])+'\n') 
#    wf2.close()
    return word_dict_sorted


#segment(readDocument())
def WCcreate(doc_path=doc_path_default,background_color=setBackgroundColor(), #背景颜色
    max_words=2000,# 词云显示的最大词数
    font_path=setFontPath(),
    mask_path=img_mask_path,#设置背景图片
    max_font_size=100, #字体最大值
    random_state=42,
    relative_scaling=0.4,
    colormap=setColorMap(),
    ):
    '''
        制作词云
        设置词云参数
    '''
    doc = readDocument(doc_path)
    segment_list = segment(doc)
    seg_list = removeStopWords(segment_list)
    mask = imread(mask_path)
    color_mask = mask # 读取背景图片,注意路径
    wc = WordCloud(
        #设置字体，不指定就会出现乱码，注意字体路径
        font_path=font_path,
        #font_path=path.join(d,'simsun.ttc'),
        #设置背景色
        background_color=background_color,
        #词云形状
        mask=color_mask,
        #允许最大词汇
        max_words=2000,
        #最大号字体
        max_font_size=60,
        colormap=colormap
    )
    wc.generate(seg_list) # 产生词云
    image_colors = ImageColorGenerator(color_mask)
    save_path = os.path.splitext(mask_path)
    save_path = save_path[0] + '_wc' + save_path[1]
    wc.to_file(save_path) #保存图片
    #  显示词云图片
    # plt.imshow(wc, interpolation="bilinear")
    # plt.axis('off')

    # #这里主要为了实现词云图片按照图片颜色取色
    # plt.figure()
    # plt.imshow(wc.recolor(color_func=image_colors), interpolation="bilinear")
    # plt.axis("off")

    # plt.show()
    return save_path


def WC_app_create(seg_list=setSegList(),background_color=setBackgroundColor(), #背景颜色
    max_words=2000,# 词云显示的最大词数
    font_path=setFontPath(),
    mask_path=img_mask_path,#设置背景图片
    max_font_size=100, #字体最大值
    random_state=42,
    relative_scaling=0.4,
    colormap=setColorMap(),
    ):
    '''
        制作词云
        设置词云参数
    '''
    mask = imread(mask_path)
    color_mask = mask # 读取背景图片,注意路径
    wc = WordCloud(
        #设置字体，不指定就会出现乱码，注意字体路径
        font_path=font_path,
        #font_path=path.join(d,'simsun.ttc'),
        #设置背景色
        background_color=background_color,
        #词云形状
        mask=color_mask,
        #允许最大词汇
        max_words=2000,
        #最大号字体
        max_font_size=60,
        colormap=colormap
    )
    wc.fit_words(seg_list) # 产生词云
    image_colors = ImageColorGenerator(color_mask)
    save_path = os.path.splitext(mask_path)
    save_path = save_path[0] + '_wc' + save_path[1]
    wc.to_file(save_path) #保存图片

    return save_path


def changeFre(word_dict,changelist):   #函数更新
    #with open('词频统计(去停用词).txt','w') as wf2:
    for item in changelist:
            if item[0] !=' ':
                    print(type(item[0]),type(item[1]))
                    word_dict[item[0]] = item[1]
    lis=[]
    lis=sorted(word_dict.items(), \
    key = lambda item:item[1], reverse=True)
    #    for key in lis:#如果更改词频，考虑从此处入手
    #        wf2.write(key[0]+' '+str(key[1])+'\n')
    #wf2.close()
    return word_dict


if __name__ == "__main__":
    doc = readDocument()
    segment_list = segment(doc)
    segment_list_remove_stopwords = removeStopWords(segment_list)
    WCcreate(segment_list_remove_stopwords)
